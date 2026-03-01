"""
Prescription Document Generator - Táº¡o ÄÆ¡n Thuá»‘c
================================================
Tool táº¡o file Word (.docx) Ä‘Æ¡n thuá»‘c theo máº«u chuáº©n bá»‡nh viá»‡n.
Dá»±a trÃªn máº«u Ä‘Æ¡n thuá»‘c cá»§a BVÄK TP Cáº§n ThÆ¡.
Khá»• giáº¥y: A5 (148mm x 210mm)

Sá»­ dá»¥ng: 
    python generate_prescription.py                    # Táº¡o tá»« file sample_data.json
    python generate_prescription.py --all              # Táº¡o táº¥t cáº£ Ä‘Æ¡n thuá»‘c
    python generate_prescription.py --id 1             # Táº¡o Ä‘Æ¡n thuá»‘c theo ID
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os
import json
import argparse

# ============================================
# Cáº¤U HÃŒNH FONT VÃ€ KÃCH THÆ¯á»šC
# ============================================
FONT_NAME = "Times New Roman"

FONT_SIZE = {
    "ministry": Pt(10),
    "hospital": Pt(10),
    "phone": Pt(9),
    "title": Pt(16),
    "barcode_text": Pt(9),
    "label": Pt(11),
    "value": Pt(11),
    "diagnosis_label": Pt(11),
    "diagnosis": Pt(11),
    "table_header": Pt(11),
    "group_name": Pt(11),
    "drug_name": Pt(12),
    "drug_instruction": Pt(11),
    "quantity": Pt(12),
    "unit": Pt(11),
    "stt": Pt(11),
    "footer": Pt(10),
    "doctor_name": Pt(11),
    "note": Pt(9),
}


def set_cell_border(cell, **kwargs):
    """Thiáº¿t láº­p viá»n cho Ã´"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ["val", "color", "sz", "space"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcBorders.append(element)
            
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement('w:{}'.format(edge))
        node.set(qn('w:w'), str(val*20))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    
    tcPr.append(tcMar)


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.0):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(before)
    paragraph_format.space_after = Pt(after)
    paragraph_format.line_spacing = line_spacing


def add_run_with_font(paragraph, text, font_size, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic
    return run


def _draw_header(doc, data):
    header_table = doc.add_table(rows=1, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    
    # Left Cell: Ministry & Hospital Info
    left_cell = header_table.cell(0, 0)
    left_cell.width = Mm(50) # TÄƒng nháº¹ width Ä‘á»ƒ chá»¯ khÃ´ng bá»‹ rá»›t dÃ²ng
    p = left_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER # CÄƒn giá»¯a khá»‘i text bÃªn trÃ¡i
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
    
    # DÃ²ng 1: Bá»˜ Y Táº¾ (ThÆ°á»ng, in hoa)
    add_run_with_font(p, "Bá»˜ Y Táº¾", FONT_SIZE["ministry"], bold=False)
    p.add_run("\n")
    
    # DÃ²ng 2: BVÄK TW Cáº¦N THÆ  (Äáº­m, Gáº¡ch chÃ¢n, in hoa)
    run = p.add_run("BVÄK TW Cáº¦N THÆ ")
    run.font.name = "Times New Roman"
    run.font.size = FONT_SIZE["hospital"]
    run.font.bold = True
    run.font.underline = True
    p.add_run("\n")
    
    # DÃ²ng 3: SÄT (Hardcode)
    # User yÃªu cáº§u thu nhá» Ä‘á»ƒ vá»«a váº·n trÃªn 1 dÃ²ng
    add_run_with_font(p, "ÄIá»†N THOáº I: 0292.382.0071", Pt(8))
    
    center_cell = header_table.cell(0, 1)
    center_cell.width = Mm(48)
    p = center_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=1, after=1, line_spacing=1.2)
    add_run_with_font(p, "ÄÆ N THUá»C", FONT_SIZE["title"], bold=True)
    center_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    right_cell = header_table.cell(0, 2)
    right_cell.width = Mm(40)
    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
    
    run = p.add_run("â•‘â•‘â•‘â•‘â•‘â•‘â•‘â•‘â•‘â•‘â•‘")
    run.font.size = Pt(12)
    p.add_run("\n")
    add_run_with_font(p, data["prescription_code"], FONT_SIZE["barcode_text"], bold=True)
    
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line_spacing=0.01) # Spacer ~6pt


def _draw_patient_info(doc, data):
    # Layout: Flow Name - Age - Gender in one line to ensure equal spacing
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=3, line_spacing=1.15)
    
    # 1. Name
    add_run_with_font(p, "Há» tÃªn:   ", Pt(10)) 
    add_run_with_font(p, data["patient"]["name"], Pt(10), bold=True)
    
    # Spacer 1
    add_run_with_font(p, "          ", Pt(10)) # Adjust spaces to control "d"
    
    # 2. Age
    add_run_with_font(p, "Tuá»•i: ", Pt(10))
    add_run_with_font(p, str(data["patient"]["age"]), Pt(10), bold=True)
    
    # Spacer 2 (Identical to Spacer 1)
    add_run_with_font(p, "          ", Pt(10))
    
    # 3. Gender
    add_run_with_font(p, "Giá»›i tÃ­nh: ", Pt(10))
    add_run_with_font(p, data["patient"]["gender"], Pt(10), bold=True)
    
    # REMOVED SPACER PARAGRAPH HERE to merge spacing visually

    # Báº£ng thÃ´ng tin chi tiáº¿t (BHYT, Äá»‹a chá»‰, Cháº©n Ä‘oÃ¡n) dÃ¹ng table Ä‘á»ƒ canh lá» Ä‘áº¹p
    info_table = doc.add_table(rows=3, cols=2)
    info_table.autofit = False
    info_table.allow_autofit = False
    
    # Cáº¥u hÃ¬nh chiá»u rá»™ng: Cá»™t nhÃ£n cá»‘ Ä‘á»‹nh, Cá»™t giÃ¡ trá»‹ láº¥y pháº§n cÃ²n láº¡i
    # Tá»•ng rá»™ng A5 margins = 148 - 10 - 10 = 128mm
    LABEL_WIDTH = Mm(30)
    VALUE_WIDTH = Mm(98)
    
    info_table.columns[0].width = LABEL_WIDTH
    info_table.columns[1].width = VALUE_WIDTH
    
    # Dá»¯ liá»‡u cho 3 dÃ²ng
    rows_data = [
        ("MÃ£ sá»‘ tháº» BHYT:", data["patient"]["insurance_code"], False, True), # (Label, Value, LabelBold, ValueBold)
        ("Äá»‹a chá»‰ liÃªn há»‡:", data["patient"]["address"], False, False),      # Address thÆ°á»ng khÃ´ng bold toÃ n bá»™, cÃ³ thá»ƒ chá»‰nh tÃ¹y Ã½
        ("Cháº©n ÄoÃ¡n:", data["diagnosis"], True, False)                       # Cháº©n ÄoÃ¡n label thÆ°á»ng Bold
    ]
    
    # Render tá»«ng dÃ²ng
    for idx, (label, value, label_bold, value_bold) in enumerate(rows_data):
        row = info_table.rows[idx]
        
        # Cell 1: Label
        c1 = row.cells[0]
        c1.width = LABEL_WIDTH
        c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_border(c1, top={'val': 'nil'}, bottom={'val': 'nil'}, left={'val': 'nil'}, right={'val': 'nil'})
        set_cell_margins(c1, top=0, bottom=0, left=0, right=0)
        
        p = c1.paragraphs[0]
        set_paragraph_spacing(p, before=3, after=3, line_spacing=1.15)
        add_run_with_font(p, label, FONT_SIZE["label"] if label != "Cháº©n ÄoÃ¡n:" else FONT_SIZE["diagnosis_label"], bold=label_bold)
        
        # Cell 2: Value
        c2 = row.cells[1]
        c2.width = VALUE_WIDTH
        c2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_border(c2, top={'val': 'nil'}, bottom={'val': 'nil'}, left={'val': 'nil'}, right={'val': 'nil'})
        set_cell_margins(c2, top=0, bottom=0, left=0, right=0)
        
        p = c2.paragraphs[0]
        set_paragraph_spacing(p, before=3, after=3, line_spacing=1.15)
        add_run_with_font(p, value, FONT_SIZE["value"] if label != "Cháº©n ÄoÃ¡n:" else FONT_SIZE["diagnosis"], bold=value_bold)

    # REMOVED SPACER PARAGRAPH AFTER INFO TABLE


def _draw_med_table_chunk(doc, data, medication_chunk, start_stt):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=2, line_spacing=1.0) # Before spacing matches Header->Info spacer
    add_run_with_font(p, "Thuá»‘c Ä‘iá»u trá»‹ :", FONT_SIZE["label"], bold=True)
    
    num_drugs = len(medication_chunk)
    num_rows = 1 + 1 + num_drugs  # Header + tÃªn BS + thuá»‘c (REMOVED TOTAL ROW)
    
    table = doc.add_table(rows=num_rows, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    table.columns[0].width = Mm(13)
    table.columns[1].width = Mm(80)
    table.columns[2].width = Mm(17.5)
    table.columns[3].width = Mm(17.5)
    
    # Header row
    header_row = table.rows[0]
    header_row.height = Mm(6)
    
    cell = header_row.cells[0]
    cell.width = Mm(13)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
    add_run_with_font(p, "STT", FONT_SIZE["table_header"], bold=True)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "FFF2CC")
    set_cell_border(cell, top={'val': 'single'}, bottom={'val': 'single'}, left={'val': 'single'}, right={'val': 'single'})
    
    cell = header_row.cells[1]
    cell.merge(header_row.cells[3]) 
    cell.width = Mm(115)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
    add_run_with_font(p, "Thuá»‘c Ä‘iá»u trá»‹", FONT_SIZE["table_header"], bold=True)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "FFF2CC")
    set_cell_border(cell, top={'val': 'single'}, bottom={'val': 'single'}, left={'val': 'single'}, right={'val': 'single'})
    
    # DÃ²ng tÃªn bÃ¡c sÄ©
    group_row = table.rows[1]
    group_row.height = Mm(7)
    cell = group_row.cells[0]
    cell.merge(group_row.cells[3])
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=2, after=2, line_spacing=1.1)
    add_run_with_font(p, data["doctor"]["name"], FONT_SIZE["group_name"], bold=True)
    set_cell_shading(cell, "FFF2CC")
    set_cell_border(cell, top={'val': 'single'}, bottom={'val': 'single'}, left={'val': 'single'}, right={'val': 'single'})
    
    # CÃ¡c dÃ²ng thuá»‘c
    for idx, drug in enumerate(medication_chunk):
        drug_row = table.rows[2 + idx]
        drug_row.height = Mm(12)
        
        # Ã” STT
        cell = drug_row.cells[0]
        cell.width = Mm(10)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
        add_run_with_font(p, str(start_stt + idx), FONT_SIZE["stt"], bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border(cell, top={'val': 'single'}, bottom={'val': 'single'}, left={'val': 'single'}, right={'val': 'single'})
        
        # Ã” Thuá»‘c
        cell = drug_row.cells[1]
        cell.width = Mm(83)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
        
        # Format: Generic (Brand Dosage) Dosage -> e.g. Losartan (Nerazzu-50 50mg) 50mg
        add_run_with_font(p, f"{drug['generic_name']} ", FONT_SIZE["drug_name"], bold=False)
        add_run_with_font(p, f"({drug['brand_name']} {drug['dosage']}) ", FONT_SIZE["drug_name"], bold=True)
        add_run_with_font(p, drug['dosage'], FONT_SIZE["drug_name"], bold=False)
        
        p.add_run("\n")
        add_run_with_font(p, drug['instructions'], FONT_SIZE["drug_instruction"], italic=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border(cell, top={'val': 'single'}, bottom={'val': 'single'}, left={'val': 'single'}, right={'val': 'nil'})
        
        # Ã” Sá»‘ lÆ°á»£ng
        cell = drug_row.cells[2]
        cell.width = Mm(17.5)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
        add_run_with_font(p, f"{drug['quantity']} ", FONT_SIZE["quantity"], bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border(cell, top={'val': 'single'}, bottom={'val': 'single'}, left={'val': 'nil'}, right={'val': 'nil'})
            
        # Ã” ÄÆ¡n vá»‹
        cell = drug_row.cells[3]
        cell.width = Mm(17.5)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
        add_run_with_font(p, drug['unit'], FONT_SIZE["unit"], bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border(cell, top={'val': 'single'}, bottom={'val': 'single'}, left={'val': 'nil'}, right={'val': 'single'})

def set_row_cant_split(row):
    """KhÃ´ng cho phÃ©p hÃ ng bá»‹ ngáº¯t trang"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)

def set_cell_shading(cell, color):
    """Thiáº¿t láº­p mÃ u ná»n"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    tcPr.append(shading)

def _draw_footer(doc, data):
    # Use a SINGLE table (2x2) for Layout to ensure integrity
    # Row 0: Advice | Info/Sign
    # Row 1: Barcode | Doctor Name
    # Row 2: Last Lines (Merged) -> "KhÃ¡m láº¡i..." + "TÃªn bá»‘/máº¹..."
    
    footer_table = doc.add_table(rows=3, cols=2)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.autofit = False
    
    for row in footer_table.rows:
        set_row_cant_split(row) # VITAL: Prevent row from splitting
        for cell in row.cells:
            set_cell_border(cell, top={'val': 'nil'}, bottom={'val': 'nil'}, left={'val': 'nil'}, right={'val': 'nil'})

    # --- ROW 0: TOP FOOTER ---
    left_top = footer_table.cell(0, 0)
    left_top.width = Mm(70)
    left_top.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    
    p = left_top.paragraphs[0]
    p.paragraph_format.keep_with_next = True 
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.1)
    run = add_run_with_font(p, "(KhÃ¡m láº¡i ngÃ y : ", FONT_SIZE["footer"])
    add_run_with_font(p, f"{data['follow_up_date']}", FONT_SIZE["footer"])
    add_run_with_font(p, ")", FONT_SIZE["footer"])
    
    p = left_top.add_paragraph()
    p.paragraph_format.keep_with_next = True 
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.1)
    add_run_with_font(p, "Lá»i dáº·n:", FONT_SIZE["footer"])
    if data.get('notes'):
        p.add_run("\n")
        add_run_with_font(p, f"{data['notes']}", FONT_SIZE["footer"], italic=True)
    
    if data.get('lab_tests'):
        p = left_top.add_paragraph()
        p.paragraph_format.keep_with_next = True
        set_paragraph_spacing(p, before=0, after=0, line_spacing=1.1)
        add_run_with_font(p, data['lab_tests'], FONT_SIZE["footer"], italic=True)
    
    right_top = footer_table.cell(0, 1)
    right_top.width = Mm(58)
    right_top.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    
    p = right_top.paragraphs[0]
    p.paragraph_format.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.1)
    add_run_with_font(p, f"NgÃ y {data['prescription_date']}", FONT_SIZE["footer"], italic=True)
    
    p = right_top.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.1)
    add_run_with_font(p, "BÃ¡c sÄ© khÃ¡m bá»‡nh", FONT_SIZE["footer"], bold=True)
    
    p = right_top.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.1)
    add_run_with_font(p, "(KÃ½, ghi rÃµ há» tÃªn)", FONT_SIZE["footer"], italic=True)
    
    for _ in range(4):
        p = right_top.add_paragraph()
        p.paragraph_format.keep_with_next = True
        set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)

    # --- ROW 1: BOTTOM FOOTER ---
    left_bottom = footer_table.cell(1, 0)
    left_bottom.width = Mm(70)
    left_bottom.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    
    p = left_bottom.paragraphs[0]
    p.paragraph_format.keep_with_next = True # Glue to next row (Last Lines)
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
    run = p.add_run("â•‘â•‘â•‘â•‘â•‘â•‘â•‘â•‘â•‘â•‘â•‘â•‘â•‘") 
    run.font.size = Pt(14)
    
    p = left_bottom.add_paragraph()
    p.paragraph_format.keep_with_next = True
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
    add_run_with_font(p, data['barcode_bottom'], FONT_SIZE["footer"])
    
    right_bottom = footer_table.cell(1, 1)
    right_bottom.width = Mm(58)
    right_bottom.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    
    p = right_bottom.paragraphs[0]
    p.paragraph_format.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
    add_run_with_font(p, f"{data['doctor']['title']}.{data['doctor']['name']}", FONT_SIZE["doctor_name"], bold=True)
    
    # --- ROW 2: LAST LINES (Merged) ---
    last_row = footer_table.rows[2]
    last_cell = last_row.cells[0]
    last_cell.merge(last_row.cells[1])
    
    p = last_cell.paragraphs[0]
    set_paragraph_spacing(p, before=1, after=0, line_spacing=1.0)
    add_run_with_font(p, "- KhÃ¡m láº¡i xin mang theo Ä‘Æ¡n nÃ y.", FONT_SIZE["note"], italic=False)
    
    p = last_cell.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
    add_run_with_font(p, "- TÃªn bá»‘ / máº¹ cá»§a tráº» hoáº·c ngÆ°á»i Ä‘Æ°a tráº» Ä‘áº¿n khÃ¡m bá»‡nh, chá»¯a bá»‡nh: ..............................", FONT_SIZE["note"], italic=False)


def calculate_layout(total_items):
    """
    TÃ­nh toÃ¡n cÃ¡ch chia thuá»‘c vÃ o cÃ¡c trang sao cho cÃ¢n Ä‘á»‘i vÃ  Ä‘áº¹p máº¯t.
    Quy táº¯c:
    - Trang thÆ°á»ng: Tá»‘i Ä‘a 8 thuá»‘c (MAX_FULL_PAGE)
    - Trang cuá»‘i: Tá»‘i Ä‘a 5 thuá»‘c (MAX_LAST_PAGE) Ä‘á»ƒ trá»« hao cho Footer.
    - Cá»‘ gáº¯ng chia Ä‘á»u thuá»‘c giá»¯a cÃ¡c trang thay vÃ¬ dá»“n háº¿t vÃ o trang Ä‘áº§u.
    """
    MAX_FULL_PAGE = 8
    MAX_LAST_PAGE = 5
    
    if total_items == 0:
        return [0]
        
    # 1. TÃ¬m sá»‘ trang tá»‘i thiá»ƒu (k)
    # CÃ´ng suáº¥t = (k-1)*8 + 5
    # TÃ¬m k nhá» nháº¥t sao cho capacity >= total_items
    k = 1
    while True:
        capacity = (k - 1) * MAX_FULL_PAGE + MAX_LAST_PAGE
        if capacity >= total_items:
            break
        k += 1
        
    if k == 1:
        return [total_items]
        
    # 2. PhÃ¢n phá»‘i Ä‘á»u thuá»‘c vÃ o k trang
    # Logic: Chia Ä‘á»u (base), pháº§n dÆ° ráº£i vÃ o cÃ¡c trang Ä‘áº§u
    base = total_items // k
    remainder = total_items % k
    
    layout = []
    for i in range(k):
        count = base + (1 if i < remainder else 0)
        layout.append(count)
        
    # 3. Kiá»ƒm tra rÃ ng buá»™c trang cuá»‘i (<= 5)
    # Náº¿u trang cuá»‘i > 5, dá»“n bá»›t item sang cÃ¡c trang trÆ°á»›c (náº¿u cÃ²n chá»— < 8)
    # Vá»›i thuáº­t toÃ¡n tÃ¬m k á»Ÿ bÆ°á»›c 1, viá»‡c nÃ y luÃ´n KHáº¢ THI.
    
    while layout[-1] > MAX_LAST_PAGE:
        diff = layout[-1] - MAX_LAST_PAGE
        # Duyá»‡t ngÆ°á»£c tá»« Ã¡p chÃ³t vá» Ä‘áº§u Ä‘á»ƒ tÃ¬m chá»— trá»‘ng
        moved = False
        for i in range(k - 2, -1, -1):
            space = MAX_FULL_PAGE - layout[i]
            if space > 0:
                move_amount = min(diff, space)
                layout[i] += move_amount
                layout[-1] -= move_amount
                diff -= move_amount
                moved = True
                if diff == 0:
                    break
        
        if not moved or diff > 0:
            # Should not happen logic wise if k is correct
            break
            
    return layout


def add_prescription_to_doc(doc, data, is_first=False):
    """
    ThÃªm 1 Ä‘Æ¡n thuá»‘c vÃ o document.
    Sá»­ dá»¥ng laytout Ä‘á»™ng Ä‘á»ƒ chia trang.
    """
    
    all_meds = data["medications"]
    total_meds = len(all_meds)
    
    # TÃ­nh toÃ¡n layout (VD: [5, 5, 5] cho 15 thuá»‘c)
    layout = calculate_layout(total_meds)
    
    # Cáº¯t thuá»‘c theo layout
    chunks = []
    start = 0
    for count in layout:
        chunks.append(all_meds[start : start + count])
        start += count
    
    if not chunks: 
        chunks = [[]]

    current_stt = 1
    for i, chunk in enumerate(chunks):
        # 1. Page Break (Náº¿u khÃ´ng pháº£i trang Ä‘áº§u tiÃªn cá»§a file lá»›n, HOáº¶C lÃ  trang tiáº¿p theo cá»§a cÃ¹ng 1 Ä‘Æ¡n thuá»‘c)
        if (not is_first) or (i > 0):
             doc.add_page_break()
             
        # Setup Page Layout (A5)
        if (not is_first) or (i > 0):
             pass

        # 2. Draw Header & Patient (Repeated on every page)
        _draw_header(doc, data)
        _draw_patient_info(doc, data)
        
        # 3. Draw Medications (Chunk)
        # start_stt truyá»n vÃ o pháº£i Ä‘Æ°á»£c tÃ­nh cá»™ng dá»“n
        _draw_med_table_chunk(doc, data, chunk, current_stt)
        current_stt += len(chunk)
        
        # 4. Draw Footer (Only on LAST chunk)
        if i == len(chunks) - 1:
            _draw_footer(doc, data)
        else:
             p = doc.add_paragraph()
             p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
             set_paragraph_spacing(p, before=1, after=1)
             add_run_with_font(p, "(Xem tiáº¿p trang sau...)", Pt(10), italic=True)


def create_prescription_doc(data_list, output_path="output/prescription.docx"):
    doc = Document()
    
    # Page Setup
    section = doc.sections[0]
    section.page_width = Mm(148) # A5
    section.page_height = Mm(210)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    
    for i, data in enumerate(data_list):
        print(f"   âœ… ÄÆ¡n thuá»‘c #{data['id']}")
        add_prescription_to_doc(doc, data, is_first=(i==0))
        
    # Ensure dir exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"\nğŸ“‚ File: {output_path}")
    print("ğŸ“„ Sá»‘ trang: {}".format(len(data_list))) # Approximation
    print("ğŸ“„ Khá»• giáº¥y: A5 (148mm x 210mm)")
    print(f"ğŸ”¤ Font: {FONT_NAME}")


def main():
    parser = argparse.ArgumentParser(description="Táº¡o Ä‘Æ¡n thuá»‘c (Layout Cáº§n ThÆ¡)")
    parser.add_argument("--data", default="generated_sample_data.json", help="File dá»¯ liá»‡u JSON input")
    parser.add_argument("--output", default="output/prescription.docx", help="ÄÆ°á»ng dáº«n file output")
    parser.add_argument("--all", action="store_true", help="Táº¡o táº¥t cáº£ Ä‘Æ¡n thuá»‘c trong data")
    parser.add_argument("--start-id", type=int, help="ID báº¯t Ä‘áº§u")
    parser.add_argument("--end-id", type=int, help="ID káº¿t thÃºc")
    
    args = parser.parse_args()
    
    try:
        print("==================================================")
        print("  ÄÆ N THUá»C GENERATOR - BVÄK TW Cáº¦N THÆ ")
        print("==================================================\n")
        
        print(f"ğŸ“‚ Äá»c dá»¯ liá»‡u tá»«: {os.path.abspath(args.data)}")
        with open(args.data, "r", encoding="utf-8") as f:
            full_data = json.load(f)
            
        all_prescriptions = full_data.get("prescriptions", [])
        
        target_list = []
        
        # Filtering logic
        if args.all:
            target_list = all_prescriptions
        elif args.start_id is not None and args.end_id is not None:
             target_list = [p for p in all_prescriptions if args.start_id <= p["id"] <= args.end_id]
             print(f"ğŸ” Lá»c Ä‘Æ¡n thuá»‘c tá»« ID {args.start_id} Ä‘áº¿n {args.end_id}")
        else:
             target_list = all_prescriptions[:1] # Default 1st
             print("âš ï¸  Cháº¿ Ä‘á»™ máº·c Ä‘á»‹nh: Chá»‰ táº¡o 1 Ä‘Æ¡n Ä‘áº§u tiÃªn. DÃ¹ng --all Ä‘á»ƒ táº¡o háº¿t.")

        if not target_list:
            print("âŒ KhÃ´ng tÃ¬m tháº¥y Ä‘Æ¡n thuá»‘c nÃ o thá»a mÃ£n Ä‘iá»u kiá»‡n!")
            return

        print(f"ğŸ“ Äang táº¡o {len(target_list)} Ä‘Æ¡n thuá»‘c vÃ o 1 file...")
        create_prescription_doc(target_list, args.output)
        
        print("\n==================================================")
        print("  HOÃ€N THÃ€NH!")
        print("==================================================")

    except Exception as e:
        print(f"\nâŒ Lá»–I: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
